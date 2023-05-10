from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text  import WD_PARAGRAPH_ALIGNMENT
from docx.image.image import *
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.shape import WD_INLINE_SHAPE
from docx2pdf import convert
from docx.shared import RGBColor
import os
from PyQt5.QtWidgets import QMessageBox






##Code here to handle patient data and selected graphs, add to functions
def get_graph_data(name, address, age, contact, graph_selected, RKneediff, LKneediff, RElbowdiff, LElbowdiff, Pelvisdiff):
    document = Document('existing_template.docx')
    name = str(name)
    address = str(address)
    age = str(age)
    contact = str(contact)
    

    p = document.add_paragraph('Name: ' + name + ' \n Age: '+ age +
                               '\n Address: ' + address + '\n Contact Num: '+ contact)
    first_run = p.runs[0]
    font = first_run.font
    font.color.rgb = RGBColor(255, 255, 255)  # white color
    font.size = Pt(16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p.space_before = Inches(3)
    p.right_indent = Inches(7) - (p.paragraph_format.left_indent or Inches(0)) - (p.paragraph_format.first_line_indent or Inches(0)) - (p.paragraph_format.space_before or Pt(0))


    description = document.add_paragraph()
    description.add_run("\n")
    description.add_run("\n")
    description.add_run("\n")
    description.add_run("\n")
    description.add_run("\n")
    description.add_run("\n")
    description.add_run("\n")
    description.add_run("\n")
    description.add_run("Pre - Post treatment difference data: (flexion difference)").bold = True
    description.add_run('\n\n               Right Knee Difference: ' + str(RKneediff) +
                        '\n\n           Left Knee Difference: ' + str(LKneediff) +
                        '\n\n           Right Elbow Difference: ' + str(RElbowdiff) +
                        '\n\n           Left Elbow Difference: ' + str(LElbowdiff) + 
                        '\n\n           Pelvis Difference: ' + str(Pelvisdiff)).bold = True
    description.alignment = WD_ALIGN_PARAGRAPH.LEFT



    for graph_name, selected in graph_selected.items():
        if selected:
            image_path = graph_name
            width = Inches(6)
            height = Inches(5)
            image = document.add_picture(image_path)
            left = Inches(2)
            top = Inches(4)
            image.left = left
            image.top = top
            image.width = width
            image.height = height
    


    # save the document
    document.save('Modified.docx')


    input_file = "Modified.docx"
    output = "report.pdf"

    convert(input_file, output)
    file_location()


def file_location():
    for root, dirs, files in os.walk("."):
        file_location = os.path.join(root, "report.pdf")
    QMessageBox.information(None, "PDF Generated", "Patient report generated successfully. Location:" + file_location, QMessageBox.Ok )




